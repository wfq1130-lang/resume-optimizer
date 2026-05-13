"""解析简历文件：支持 PDF、DOCX、TXT"""

import os


def parse_resume(filepath, filename):
    """根据文件扩展名调用对应解析器，返回纯文本"""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return _parse_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return _parse_docx(filepath)
    elif ext in (".txt", ".md", ".markdown"):
        return _parse_txt(filepath)
    else:
        # 尝试作为文本读取
        try:
            return _parse_txt(filepath)
        except Exception:
            return ""


def _parse_pdf(filepath):
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(filepath)
        text_parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
        return "\n".join(text_parts)
    except ImportError:
        return "[错误] 未安装 PyPDF2 库"
    except Exception as e:
        return f"[错误] PDF解析失败: {str(e)}"


def _parse_docx(filepath):
    try:
        from docx import Document
        doc = Document(filepath)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)
    except ImportError:
        return "[错误] 未安装 python-docx 库"
    except Exception as e:
        return f"[错误] DOCX解析失败: {str(e)}"


def _parse_txt(filepath):
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    for enc in encodings:
        try:
            with open(filepath, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    return "[错误] 无法识别文件编码"
